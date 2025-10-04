from typing import Annotated, Any, Optional
from fastapi import UploadFile, File, HTTPException
from fastapi.responses import JSONResponse


from datetime import datetime, timezone
from fastapi import HTTPException
from fastapi.responses import JSONResponse
import os
import tempfile
import PyPDF2
import docx
from doc2docx import convert as doc_to_docx_convert
import magic
import hashlib
import re
import asyncio
from pathlib import Path
import logging


from typing import Optional, AsyncGenerator
import tempfile
import contextlib
from pathlib import Path


from pydantic import BaseModel


from app.utils.common.get_cid import get_cid
from app.configs import Configs


CID_PATTERN = re.compile(r'^bafkreiht[a-z0-9]{52}$')
CID_LENGTH = 60  # Length of 'bafkreiht' + 52 characters

def validate_client_cid(cid: Optional[str] = None) -> None:
    """Validate client CID format."""
    if cid is None:
        return

    if not isinstance(cid, str):
        raise TypeError(f"client_cid must be a string or None, got {type(cid).__name__}")

    if not cid.strip():  # Empty string
        raise ValueError("Client CID cannot be empty or whitespace")

    if len(cid) != CID_LENGTH:  # bafkreiht(8) + 52 chars
        raise ValueError(f"Client CID must be exactly 60 characters, got {len(cid)}")

    if not cid.startswith('bafkreiht'):
        raise ValueError("Client CID must start with 'bafkreiht'")

    # Check for invalid characters (unicode, spaces, etc)
    if not re.match(CID_PATTERN, cid):
        raise ValueError("Client CID contains invalid characters")

def validate_file(file: UploadFile) -> None:
    pass


def _match_mime(ext: str, mime: str, expected_mimes: dict[str, list[str]]) -> None:
    """Helper to match MIME types."""
    expected_mime = expected_mimes.get(ext)
    if mime not in expected_mime:
        raise HTTPException(
            status_code=400, 
            detail=f"{ext.upper()} file appears to be invalid or corrupt. Expected {expected_mime}, got {mime}"
        )


class UploadDocumentModel(BaseModel):
    file: Annotated[UploadFile, validate_file]
    client_cid: Optional[Annotated[str, validate_client_cid]] = None


@contextlib.asynccontextmanager
async def ram_temp_file(suffix: str) -> AsyncGenerator[Path, None]:
    """Context manager for temporary files in RAM."""
    # Use /dev/shm for Linux (RAM), falls back to system temp
    temp_dir = Path("/dev/shm") if Path("/dev/shm").exists() else None
    print(f"DEBUG temp_dir: {temp_dir}")

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False, dir=temp_dir) as tmp:

        tmp_path = Path(tmp.name)

        try:
            yield tmp_path
        finally:
            try:
                tmp_path.unlink()
            except FileNotFoundError:
                pass


async def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    text_content = []
    idx = None
    try:
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            if num_pages == 0:
                return ""

            for idx in range(num_pages):
                text = pdf_reader.pages[idx].extract_text()
                text_content.append(text)

        return '\n'.join(text_content).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from page {idx+1} of PDF '{file_path}': {e}") from e


async def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    text_content = []
    idx = None
    try:
        doc = docx.Document(file_path)
        for paragraph in enumerate(doc.paragraphs, start=1):
            text_content.append(paragraph.text)
        return '\n'.join(text_content).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from paragraph {idx} of DOCX '{file_path}': {e}") from e


async def _extract_text_from_doc(file_path: str) -> str:
    """Extract text from legacy DOC file."""
    text = ""
    try:
        # Convert DOC to DOCX first
        async with ram_temp_file(suffix='docx') as docx_path:
            doc_to_docx_convert(file_path, docx_path)
            text = await _extract_text_from_docx(docx_path)
            return text
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOC: {str(e)}")


async def _extract_text_from_txt(
    file_path: str, 
    encodings: set[str] = {'utf-8', 'latin-1', 'cp1252', 'iso-8859-1'}
    ) -> str:
    """Extract text from plain text file."""
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as txt_file:
                return txt_file.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise ValueError(f"Failed to read text file '{file_path}': {e}") from e
    else:
        raise ValueError(f"Failed to decode text file '{file_path}' with encodings: {', '.join(encodings)}")


class UploadDocument:

    def __init__(self, 
                 *, 
                 resources: dict[str, Any] = None, 
                 configs: Configs = None
                 ) -> None:

        self.configs = configs
        self.resources = resources

        self.logger: logging.Logger = self.resources["logger"]

        self.supported_file_types: set[str] = self.configs.SUPPORTED_FILE_TYPES
        self.max_file_size: int = self.configs.MAX_FILE_SIZE_BYTES

        # TODO Move these to configs object
        self.expected_mimes: dict[str, list[str]] = {
            'pdf': ['application/pdf'],
            'doc': ['application/msword'],
            'docx': [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                "application/zip",
                "application/octet-stream"
            ],
            'txt': ['text/plain']
        }
        self.allowed_text_mimes = ['application/octet-stream', 'application/x-empty']


    async def upload_document(self, file: UploadFile, temp_storage: dict, client_cid: Optional[str] = None) -> JSONResponse:
        """
        API endpoint to handle document uploads from the frontend.

        This method processes document uploads from the JavaScript frontend's uploadDocument
        function. It integrates with the existing document processing pipeline to:
        - Validate uploaded file format and size constraints
        - Extract text content from PDF, DOC, DOCX, TXT files.
        - Generate content identifiers (CIDs) for document tracking
        - Store document metadata and content in the database
        - Optionally associate uploads with client sessions for history tracking
        """
        # Initialize timestamp at the start
        timestamp = datetime.now(timezone.utc).isoformat().replace('+00:00', 'Z')

        if not isinstance(temp_storage, dict):
            raise TypeError(f"temp_storage must be a dict, got {type(temp_storage).__name__}")

        if not isinstance(file, UploadFile):
            raise TypeError(f"file must be an UploadFile object, got {type(file).__name__}")

        if client_cid is not None and not isinstance(client_cid, str):
            raise TypeError(f"client_cid must be a string or None, got {type(client_cid).__name__}")

        try:
            validate_client_cid(client_cid)
        except (TypeError, ValueError) as e:
            raise

        # Main processing logic
        try:
            # Get file details
            filename: str = file.filename or ""
            if not filename:
                raise ValueError("file.filename was empty")

            # Validate file has an extension
            if '.' not in filename:
                raise HTTPException(status_code=415, detail="File must have an extension")

            ext = filename.split('.')[-1].lower().strip()

            if ext not in self.supported_file_types:
                raise HTTPException(status_code=415, detail=f"Unsupported file type: {ext}")

            # Read file content
            try:
                file_content = await file.read()
                await file.seek(0)  # Reset file pointer
            except Exception as e:
                raise IOError(f"Failed to read file content: {e}") from e

            # Check file size
            print(f"file_content for {ext}:{file_content[:100]}...")
            file_size = len(file_content)

            # Check for empty file
            if file_size <= 0:
                return JSONResponse(
                    status_code=200,
                    content={
                        "status": "error",
                        "message": "File is empty",
                        "error_code": "FILE_EMPTY",
                        "upload_timestamp": timestamp
                    }
                )
            elif file_size > self.max_file_size:
                raise HTTPException(
                    status_code=413,
                    detail=f"File size {file_size} exceeds maximum allowed size of {self.max_file_size} bytes"
                )

            # Verify file integrity (check if it's actually the claimed type)
            try:
                mime = magic.from_buffer(file_content, mime=True)
                self.logger.debug(f"Detected MIME type: {mime} for file: {filename}")

                # For text files, be more lenient with MIME type
                match ext:
                    case ext if ext in self.supported_file_types:
                        expected_mime = self.expected_mimes.get(ext)
                        if mime not in expected_mime:
                            raise HTTPException(
                                status_code=400, 
                                detail=f"{ext.upper()} file appears to be corrupted or invalid. Expected {expected_mime}, got {mime}"
                            )
                    case _:
                        raise ValueError(f"Unsupported file type for text extraction: {ext}")
            except HTTPException:
                raise
            except ValueError:
                raise
            except Exception as e:
                raise RuntimeError(f"Unexpected error determining file mimetype: {e}") from e

            # Generate CID from content
            cid = get_cid(file_content)

            # Save file temporarily for text extraction
            async with ram_temp_file(ext) as tmp_file:
                tmp_file.write_bytes(file_content)
                tmp_file_path = str(tmp_file)

                # Extract text based on file type
                match ext:
                    case 'pdf':
                        extracted_text = await _extract_text_from_pdf(tmp_file_path)
                    case 'docx':
                        extracted_text = await _extract_text_from_docx(tmp_file_path)
                    case 'doc':
                        extracted_text = await _extract_text_from_doc(tmp_file_path)
                    case 'txt':
                        extracted_text = await _extract_text_from_txt(tmp_file_path)
                    case _:
                        raise ValueError(f"Unsupported file type for text extraction: {ext}")

            # Store metadata and content in temp_storage dict
            insert_dict = {
                "filename": filename,
                "file_size": file_size,
                "mime": mime,
                "extracted_text": extracted_text,
                "client_cid": client_cid,
                "timestamp": timestamp,
                "file_content": file_content
            }
            temp_storage[cid] = insert_dict

            # Return success response
            return JSONResponse(
                status_code=200,
                content={
                    "status": "success",
                    "message": "File uploaded successfully",
                    "cid": cid,
                    "filename": filename,
                    "file_size": file_size,
                    "upload_timestamp": timestamp
                }
            )

        except HTTPException as e:
            # Re-raise HTTP exceptions as-is
            raise
        except TypeError as e:
            # Re-raise type errors for parameter validation
            raise
        except ValueError as e:
            # Re-raise value errors for validation
            raise
        except IOError as e:
            # Re-raise IO errors
            raise
        except Exception as e:
            # Catch any other unexpected errors
            self.logger.exception(f"Unexpected error in upload_document: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Internal server error: {str(e)}"
            )

def make_upload_document():
    """Factory function to create UploadDocument instance with default resources and configs."""
    from app.configs import configs
    from app.logger import logger

    resources = {
        "logger": logger,
    }
    return UploadDocument(resources=resources, configs=configs)
